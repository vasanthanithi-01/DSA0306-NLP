from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from pathlib import Path
from collections import Counter
import json
import re
import uuid
import math
import os
from google import genai
from google.genai import types
from dotenv import load_dotenv
import numpy as np
import faiss
from docx import Document
from PIL import Image
import pytesseract
import pymupdf
from sentence_transformers import SentenceTransformer


app = FastAPI(title="DocInsight AI API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:5173",
        "http://127.0.0.1:5173",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
load_dotenv()
gemini_client = genai.Client(
    api_key=os.getenv("GEMINI_API_KEY")
)

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR.mkdir(exist_ok=True)
DATA_DIR.mkdir(exist_ok=True)

TESSERACT_PATH = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
if TESSERACT_PATH.exists():
    pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_PATH)

print("Loading embedding model...")
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
print("Embedding model loaded.")

STOP_WORDS = {
    "the","and","for","that","with","this","from","are","was","were",
    "has","have","had","into","than","then","their","they","them","its",
    "our","your","you","not","but","all","can","will","also","using",
    "used","use","which","about","more","some","other","only","each",
    "these","those","document","page","pages","report","file","been",
    "being","shall","should","would","could","may","might","such","than",
    "there","where","when","what","while","within","through","between",
    "under","over","after","before","during","onto","upon","into","out"
}

# In-memory document store. Restarting the backend clears uploaded indexes.
DOCUMENTS = {}


def clean_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"https?://\S+|www\.\S+|localhost:\d+/?\S*", " ", text, flags=re.I)
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if re.fullmatch(r"\d+\s*/\s*\d+", line):
            continue
        if re.fullmatch(r"(page\s*)?\d+", line, flags=re.I):
            continue
        lines.append(line)
    text = "\n".join(lines)
    text = re.sub(
        r"\b\d{1,2}/\d{1,2}/\d{2,4},?\s*\d{1,2}:\d{2}\s*(AM|PM)?\b",
        " ", text, flags=re.I
    )
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def extract_pdf_pages(file_path: Path):
    try:
        pages = []
        pdf = pymupdf.open(str(file_path))
        for page_number, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if not text:
                try:
                    pix = page.get_pixmap(matrix=pymupdf.Matrix(2, 2))
                    mode = "RGBA" if pix.alpha else "RGB"
                    image = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                    text = pytesseract.image_to_string(image)
                except Exception:
                    text = ""
            pages.append({"page": page_number, "text": clean_text(text)})
        pdf.close()
        return pages
    except Exception as error:
        raise HTTPException(status_code=400, detail=f"Could not process PDF: {error}")


def extract_docx_pages(file_path: Path):
    try:
        document = Document(str(file_path))
        content = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append(text)
        for table in document.tables:
            for row in table.rows:
                row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_content:
                    content.append(" | ".join(row_content))
        return [{"page": 1, "text": clean_text("\n".join(content))}]
    except Exception as error:
        raise HTTPException(status_code=400, detail=f"Could not process DOCX: {error}")


def extract_image_pages(file_path: Path):
    try:
        image = Image.open(file_path)
        text = clean_text(pytesseract.image_to_string(image))
        return [{"page": 1, "text": text}]
    except Exception as error:
        raise HTTPException(status_code=400, detail=f"Could not OCR image: {error}")


def chunk_pages(pages, chunk_size=500, overlap_sentences=1):
    chunks = []

    for page_data in pages:
        page_number = page_data["page"]
        text = page_data["text"]

        if not text:
            continue

        # Split text into sentences
        sentences = re.split(
            r"(?<=[.!?])\s+",
            text.strip()
        )

        current_chunk = []
        current_length = 0
        previous_sentences = []

        for sentence in sentences:
            sentence = sentence.strip()

            if not sentence:
                continue

            sentence_length = len(sentence)

            # If adding this sentence exceeds the chunk size,
            # save the current chunk first
            if (
                current_chunk
                and current_length + sentence_length + 1 > chunk_size
            ):
                chunks.append({
                    "chunk_id": len(chunks),
                    "page": page_number,
                    "text": " ".join(current_chunk)
                })

                # Keep the last complete sentence(s) as overlap
                previous_sentences = current_chunk[
                    -overlap_sentences:
                ]

                current_chunk = previous_sentences.copy()
                current_length = len(
                    " ".join(current_chunk)
                )

            current_chunk.append(sentence)
            current_length += sentence_length + 1

        # Save the remaining chunk
        if current_chunk:
            chunks.append({
                "chunk_id": len(chunks),
                "page": page_number,
                "text": " ".join(current_chunk)
            })

    return chunks


def normalize_embeddings(vectors: np.ndarray) -> np.ndarray:
    vectors = np.asarray(vectors, dtype=np.float32)
    faiss.normalize_L2(vectors)
    return vectors


def extract_keywords(text: str, limit=12):
    words = re.findall(r"[A-Za-z][A-Za-z'-]{2,}", text.lower())
    counts = Counter(word for word in words if word not in STOP_WORDS)
    return [{"keyword": word, "count": count} for word, count in counts.most_common(limit)]


def split_sentences(text: str):
    sentences = re.split(r"(?<=[.!?])\s+|\n+", text)
    return [s.strip() for s in sentences if len(s.strip()) >= 35]


def build_summary(pages, max_sentences=4):
    text = " ".join(p["text"] for p in pages if p["text"])

    # Get candidate sentences
    raw_sentences = split_sentences(text)

    candidates = []

    for sentence in raw_sentences:
        sentence = sentence.strip()

        # Ignore fragments
        if len(sentence) < 50:
            continue

        # Ignore code / JSON / malformed content
        if any(char in sentence for char in ["{", "}", "[", "]", "=>"]):
            continue

        # Ignore sentences with too little normal language
        letters = sum(char.isalpha() for char in sentence)

        if letters / max(len(sentence), 1) < 0.6:
            continue

        # Ignore common document metadata / noise
        if re.search(
            r"\b(Name|Reg No|Bloom.?s Level|Marks|Code of Conduct)\b",
            sentence,
            re.IGNORECASE
        ):
            continue

        candidates.append(sentence)

    if not candidates:
        return "A meaningful summary could not be generated from the extracted document."

    # Prevent huge documents from slowing down processing
    candidates = candidates[:100]

    # Create embeddings for all candidate sentences
    embeddings = embedding_model.encode(
        candidates,
        convert_to_numpy=True,
        show_progress_bar=False
    )

    embeddings = normalize_embeddings(embeddings)

    # Calculate the overall meaning of the document
    document_embedding = np.mean(
        embeddings,
        axis=0,
        keepdims=True
    )

    document_embedding = normalize_embeddings(document_embedding)

    # Find sentences most semantically representative
    scores = np.dot(
        embeddings,
        document_embedding.T
    ).flatten()

    # Select more candidates than needed first
    ranked_indices = np.argsort(scores)[::-1]

    selected = []

    for index in ranked_indices:
        candidate = candidates[int(index)]

        # Avoid near-duplicate information
        if selected:
            candidate_embedding = embeddings[index]

            selected_embeddings = np.array([
                embeddings[candidates.index(sentence)]
                for sentence in selected
            ])

            similarity = np.dot(
                selected_embeddings,
                candidate_embedding
            )

            # Skip if too similar to something already selected
            if np.max(similarity) > 0.85:
                continue

        selected.append(candidate)

        if len(selected) >= max_sentences:
            break

    # Put selected sentences back into document order
    selected.sort(key=lambda sentence: candidates.index(sentence))

    return " ".join(selected)


def build_key_points(chunks, embeddings, limit=5):
    if not chunks:
        return []
    centroid = embeddings.mean(axis=0, keepdims=True)
    centroid = normalize_embeddings(centroid)
    scores = (embeddings @ centroid.T).reshape(-1)
    best = np.argsort(scores)[::-1]

    points = []
    used_pages = set()
    for idx in best:
        chunk = chunks[int(idx)]
        if chunk["page"] in used_pages and len(points) < 3:
            continue
        sentences = split_sentences(chunk["text"])
        text = sentences[0] if sentences else chunk["text"]
        points.append({
            "chunk_id": chunk["chunk_id"],
            "page": chunk["page"],
            "text": text[:450]
        })
        used_pages.add(chunk["page"])
        if len(points) >= limit:
            break
    return points


def extract_entities(text: str):
    emails = sorted(set(re.findall(r"\b[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}\b", text)))
    urls = sorted(set(re.findall(r"\bhttps?://[^\s<>()]+", text)))
    dates = sorted(set(re.findall(
        r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|"
        r"\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b",
        text, flags=re.I
    )))

    candidates = re.findall(r"\b(?:[A-Z][A-Za-z&.-]*\s+){1,4}[A-Z][A-Za-z&.-]*\b", text)
    names = []
    seen = set()
    for item in candidates:
        item = " ".join(item.split())
        if item.lower() not in seen and len(item) > 2:
            seen.add(item.lower())
            names.append(item)
        if len(names) >= 15:
            break

    return {
        "names_or_organizations": names,
        "dates": dates[:15],
        "emails": emails[:15],
        "urls": urls[:15]
    }


def retrieve(document_id: str, query: str, top_k: int = 3):
    document = DOCUMENTS.get(document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found. Please upload it again.")

    query_vector = embedding_model.encode([query], convert_to_numpy=True)
    query_vector = normalize_embeddings(query_vector)

    k = min(max(1, top_k), len(document["chunks"]))
    scores, indices = document["index"].search(query_vector, k)

    results = []
    for rank, (score, idx) in enumerate(zip(scores[0], indices[0]), start=1):
        if idx < 0:
            continue
        chunk = document["chunks"][int(idx)]
        results.append({
            "rank": rank,
            "chunk_id": chunk["chunk_id"],
            "page": chunk["page"],
            "text": chunk["text"],
            "score": float(max(0.0, min(1.0, score)))
        })
    return results


class SearchRequest(BaseModel):
    document_id: str
    query: str
    top_k: int = 3


class AskRequest(BaseModel):
    document_id: str
    question: str
    top_k: int = 3


@app.get("/")
def root():
    return {"status": "ok", "message": "DocInsight AI API is running"}


@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")

    original_name = file.filename or "uploaded_document"
    extension = Path(original_name).suffix.lower()
    supported = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}
    if extension not in supported:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format: {extension or 'unknown'}. Supported: PDF, DOCX, PNG, JPG, JPEG, TIFF, BMP."
        )

    document_id = str(uuid.uuid4())
    file_path = UPLOAD_DIR / f"{document_id}{extension}"
    file_path.write_bytes(contents)

    try:
        if extension == ".pdf":
            pages = extract_pdf_pages(file_path)
        elif extension == ".docx":
            pages = extract_docx_pages(file_path)
        else:
            pages = extract_image_pages(file_path)

        total_text = "\n".join(p["text"] for p in pages).strip()
        if not total_text:
            raise HTTPException(
                status_code=400,
                detail="Could not extract any readable text from this file."
            )

        chunks = chunk_pages(pages)
        if not chunks:
            raise HTTPException(status_code=400, detail="No text chunks could be created.")

        texts = [chunk["text"] for chunk in chunks]
        embeddings = embedding_model.encode(
            texts,
            convert_to_numpy=True,
            show_progress_bar=False
        )
        embeddings = normalize_embeddings(embeddings)

        dimension = int(embeddings.shape[1])
        index = faiss.IndexFlatIP(dimension)
        index.add(embeddings)

        summary = build_summary(pages)
        keywords = extract_keywords(total_text)
        key_points = build_key_points(chunks, embeddings)
        entities = extract_entities(total_text)

        words = len(re.findall(r"\b\w+\b", total_text))
        characters = len(total_text)
        reading_time = max(1, math.ceil(words / 200))

        DOCUMENTS[document_id] = {
            "filename": original_name,
            "pages": pages,
            "chunks": chunks,
            "embeddings": embeddings,
            "index": index,
        }

        return {
            "success": True,
            "document_id": document_id,
            "filename": original_name,
            "file_type": extension.replace(".", "").upper(),
            "pages": len(pages),
            "words": words,
            "characters": characters,
            "chunks": len(chunks),
            "reading_time_minutes": reading_time,
            "embedding_dimension": dimension,
            "summary": summary,
            "keywords": keywords,
            "key_points": key_points,
            "entities": entities,
        }

    except HTTPException:
        if file_path.exists():
            file_path.unlink()
        raise
    except Exception as error:
        if file_path.exists():
            file_path.unlink()
        raise HTTPException(status_code=500, detail=f"Document processing failed: {error}")


@app.post("/search")
def semantic_search(request: SearchRequest):
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Search query cannot be empty.")
    results = retrieve(request.document_id, request.query.strip(), request.top_k)
    return {
        "query": request.query,
        "results": results
    }




@app.post("/ask")
def ask_document(request: AskRequest):

    if not request.question.strip():
        raise HTTPException(
            status_code=400,
            detail="Question cannot be empty."
        )

    # Step 1: Retrieve relevant document chunks
    retrieved_chunks = retrieve(
        request.document_id,
        request.question.strip(),
        request.top_k
    )

    # Step 2: Combine retrieved chunks into RAG context
    context = "\n\n".join(
        f"[Page {chunk['page']}]\n{chunk['text']}"
        for chunk in retrieved_chunks
    )

    # Step 3: Ground the LLM strictly in the retrieved context
    prompt = f"""
You are DocInsight AI, an intelligent document question-answering assistant.

Answer the user's question using ONLY the provided document context.

Rules:
- Give a clear, direct, and concise answer.
- Do not invent information.
- If the answer is not present in the context, say:
  "I couldn't find that information in the uploaded document."
- Do not mention that you were given chunks or retrieved context.
- When useful, mention the relevant page number.

DOCUMENT CONTEXT:
{context}

USER QUESTION:
{request.question}

ANSWER:
"""

    # Step 4: Generate answer using Gemini
    try:
        response = gemini_client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.2,
                max_output_tokens=500
            )
        )

        answer = response.text

        if not answer:
            raise Exception("No answer generated.")

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Answer generation failed: {str(e)}"
        )

    return {
        "question": request.question,
        "answer": answer,
        "sources": [
            {
                "page": chunk["page"],
                "chunk_id": chunk["chunk_id"],
                "score": chunk["score"]
            }
            for chunk in retrieved_chunks
        ]
    }